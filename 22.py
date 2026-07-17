import os
import re
import io
import time
import datetime
from typing import List, Optional, Dict, Any, Generator
import requests
from requests.adapters import HTTPAdapter
from urllib3.util import Retry
from bs4 import BeautifulSoup
from pydantic import BaseModel

# Database (SQLite Enterprise)
from sqlalchemy import Column, Integer, String, Text, DateTime, Float, create_engine
from sqlalchemy.orm import declarative_base, sessionmaker, Session

# FastAPI Core & Security
from fastapi import FastAPI, Depends, HTTPException, status, File, UploadFile, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from passlib.context import CryptContext
from jose import jwt, JWTError

# Environment Auto Loader
from dotenv import load_dotenv
load_dotenv()

SECRET_KEY = os.getenv("JWT_SECRET", "9a6f3b1e7c5d8a2f4e0b1c3d5e7f9a0b1c2d3e4f5a6b7c8d9e0f")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
ALGORITHM = "HS256"

# Advanced Production File Parsers
import pypdf
import docx
import openpyxl

# Vector DB & Unified LangChain Hub
import chromadb
from langchain_community.vectorstores import Chroma
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ==========================================
# 1. 雲端/容器資料持久化儲存路徑優化
# ==========================================
# 自動偵測是否在 Render 等雲端環境，自動將資料定錨在持久化硬碟區
DATA_DIR = "/app/data" if os.path.exists("/app/data") or os.environ.get("RENDER") else "./data"
os.makedirs(DATA_DIR, exist_ok=True)

DATABASE_URL = f"sqlite:///{os.path.join(DATA_DIR, 'enterprise_knowledge_hub.db')}"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class DBUser(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    hashed_password = Column(String)
    role = Column(String, default="user")

class DBDocument(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, unique=True)
    source_type = Column(String)
    chunk_count = Column(Integer, default=0)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class DBChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, index=True)
    role = Column(String) 
    content = Column(Text)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)

class DBLog(Base):
    __tablename__ = "system_logs"
    id = Column(Integer, primary_key=True, index=True)
    user = Column(String, default="System")
    category = Column(String)
    message = Column(Text)
    model = Column(String, default="")
    execution_time = Column(Float, default=0.0)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)

Base.metadata.create_all(bind=engine)

def get_db():
    db = SessionLocal()
    try: yield db
    finally: db.close()

def log_event(category: str, message: str, user: str = "System", model: str = "", exec_time: float = 0.0):
    db = SessionLocal()
    db.add(DBLog(user=user, category=category, message=message, model=model, execution_time=exec_time))
    db.commit()
    db.close()

# ==========================================
# 2. 混合雲端雙核心 RAG 引擎 (支援 OpenAI / Ollama 彈性切換)
# ==========================================
class EnterpriseRAGEngine:
    def __init__(self):
        # 核心亮點：讀取環境變數自動切換模型供應商，雲端/本地通吃！
        self.provider = os.getenv("LLM_PROVIDER", "ollama").lower()
        chroma_path = os.path.join(DATA_DIR, "chroma_db")

        if self.provider == "openai":
            from langchain_openai import OpenAIEmbeddings
            self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
            self.llm_model = "gpt-4o-mini"
            print(f"🚀 [RAG 核心系統] 雲端模式啟動：使用 OpenAI ({self.llm_model})")
        else:
            self.embeddings = OllamaEmbeddings(base_url=OLLAMA_BASE_URL, model="bge-m3")
            self.llm_model = "qwen2.5:7b"
            print(f"🏠 [RAG 核心系統] 本地模式啟動：使用 Ollama ({self.llm_model})")

        # 關鍵點：使用 PersistentClient，確保雲端/重啟磁碟資料永久儲存
        self.chroma_client = chromadb.PersistentClient(path=chroma_path)
        self._init_collection()
        self.text_splitter = RecursiveCharacterCharacterTextSplitter(chunk_size=500, chunk_overlap=80)

    def _init_collection(self):
        self.vector_store = Chroma(
            client=self.chroma_client,
            collection_name="enterprise_hub_vector_store",
            embedding_function=self.embeddings
        )

    def add_content(self, text: str, source_name: str, source_type: str, author: str = "System") -> int:
        chunks = self.text_splitter.split_text(text)
        if not chunks: return 0
        
        docs = []
        for i, chunk in enumerate(chunks):
            docs.append(LCDocument(
                page_content=chunk, 
                metadata={"source": source_name, "chunk_id": i, "author": author}
            ))
        self.vector_store.add_documents(docs)
        
        db = SessionLocal()
        existing = db.query(DBDocument).filter(DBDocument.filename == source_name).first()
        if existing:
            existing.chunk_count += len(chunks)
        else:
            db.add(DBDocument(filename=source_name, source_type=source_type, chunk_count=len(chunks)))
        db.commit()
        db.close()
        return len(chunks)

    def query_stream(self, question: str, user: str = "Guest") -> Generator[str, None, None]:
        start_time = time.time()
        db = SessionLocal()
        
        histories = db.query(DBChatHistory).filter(DBChatHistory.username == user).order_by(DBChatHistory.timestamp.desc()).limit(6).all()
        memory_str = "".join([f"{'使用者' if h.role=='user' else '助手'}: {h.content}\n" for h in reversed(histories)])

        try:
            raw_docs_with_scores = self.vector_store.similarity_search_with_score(question, k=10)
        except Exception:
            raw_docs_with_scores = []

        if not raw_docs_with_scores:
            yield f"data: 目前企業知識庫中查無任何文件索引，請先在左側面板上傳內部規章。\\n\\n"
            db.close()
            return

        reranked = sorted(
            [(doc, score - (len(set(list(question)).intersection(set(list(doc.page_content)))) * 0.02)) for doc, score in raw_docs_with_scores],
            key=lambda x: x[1]
        )
        final_top_docs = [d for d, s in reranked[:4]]
        
        # 雲端/本地判定信心門檻
        threshold = 1.6 if self.provider == "openai" else 1.3
        valid_docs = [d for d in final_top_docs if any(doc == d for doc, sc in raw_docs_with_scores if sc < threshold)]

        if not valid_docs:
            log_event("AI", f"提問: '{question}' -> 信心度不足阻斷。", user=user, model=self.llm_model)
            yield f"data: 根據安全規則：知識庫中查無與此問題相關之高信心片段，無法回答內部非公開政策。\\n\\n"
            db.close()
            return

        context = "\n".join([f"【來源: {d.metadata.get('source')}】: {d.page_content}" for d in valid_docs])
        system_prompt = (
            "你是一位專業嚴謹的企業知識庫助手。請完全且僅能依據下方提供的【參考資料】進行回答。\n"
            "若資料不足以完全回答，請直接說：『知識庫沒有相關資料。』，嚴禁衍生個人外部常識或編造內容。\n\n"
            f"【歷史前情提要】:\n{memory_str}\n"
            f"【參考資料】:\n{context}\n\n"
            f"【使用者當前問題】: {question}"
        )

        total_generated_text = ""
        
        # 根據不同的引擎初始化串流
        if self.provider == "openai":
            from langchain_openai import ChatOpenAI
            llm = ChatOpenAI(model=self.llm_model, temperature=0.0, streaming=True)
        else:
            llm = ChatOllama(base_url=OLLAMA_BASE_URL, model=self.llm_model, temperature=0.0, streaming=True)

        for chunk in llm.stream(system_prompt):
            content = chunk.content
            total_generated_text += content
            yield f"data: {content.replace(chr(10), '\\n')}\n\n"
        
        source_footer = "\\n\\n---\\n### 📌 精準交叉引用來源"
        for doc in valid_docs[:2]:
            source_footer += f"\\n* 檔案來源: `{doc.metadata.get('source')}`"
        yield f"data: {source_footer}\n\n"

        db.add(DBChatHistory(username=user, role="user", content=question))
        db.add(DBChatHistory(username=user, role="assistant", content=total_generated_text))
        db.commit()
        db.close()
        
        log_event("AI", f"提問: '{question}' 成功解答。", user=user, model=self.llm_model, exec_time=time.time()-start_time)

rag_engine = EnterpriseRAGEngine()

# ==========================================
# 3. FastAPI Web 應用與高級響應式前端
# ==========================================
app = FastAPI(title="Enterprise Knowledge Hub Production v2.5")

@app.get("/status")
def get_system_status():
    return {"provider": rag_engine.provider, "model": rag_engine.llm_model, "storage": "Persistent"}

@app.get("/", response_class=HTMLResponse)
def index_page():
    # 內嵌完整且精美的響應式前端：支援 RWD 手機、平板、電腦
    return """
    <!DOCTYPE html>
    <html lang="zh-TW">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>企業級在地智能知識庫系統</title>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
        <style>
            :root {
                --primary: #0f172a; --primary-light: #1e293b; --accent: #2563eb; --accent-hover: #1d4ed8;
                --bg: #f8fafc; --panel-bg: #ffffff; --text: #334155; --text-dark: #0f172a;
                --text-light: #64748b; --border: #e2e8f0; --radius: 12px;
                --shadow: 0 4px 6px -1px rgb(0 0 0 / 0.05);
            }
            * { box-sizing: border-box; margin: 0; padding: 0; font-family: 'Inter', -apple-system, sans-serif; }
            body { background-color: var(--bg); color: var(--text); display: flex; flex-direction: column; height: 100vh; }
            header { background: var(--primary); color: white; padding: 16px 24px; display: flex; justify-content: space-between; align-items: center; box-shadow: var(--shadow); }
            header h1 { font-size: 1.25rem; font-weight: 600; display: flex; align-items: center; gap: 8px; }
            .badge { background: var(--accent); font-size: 0.75rem; padding: 4px 8px; border-radius: 20px; font-weight: 500; }
            .container { display: grid; grid-template-columns: 320px 1fr; flex: 1; overflow: hidden; height: calc(100vh - 60px); }
            .sidebar { background: var(--panel-bg); border-right: 1px solid var(--border); padding: 24px; display: flex; flex-direction: column; gap: 24px; overflow-y: auto; }
            .main-panel { display: flex; flex-direction: column; height: 100%; background: #fafafa; overflow: hidden; }
            .card { background: var(--panel-bg); border: 1px solid var(--border); border-radius: var(--radius); padding: 20px; }
            h2 { font-size: 0.85rem; text-transform: uppercase; letter-spacing: 0.05em; color: var(--text-light); margin-bottom: 12px; font-weight: 700; }
            .file-dropzone { border: 2px dashed var(--border); border-radius: var(--radius); padding: 25px 15px; text-align: center; cursor: pointer; background: var(--bg); transition: 0.2s; }
            .file-dropzone:hover { border-color: var(--accent); background: #eff6ff; }
            input[type="file"] { display: none; }
            .btn { background: var(--accent); color: white; border: none; padding: 12px 16px; border-radius: 8px; cursor: pointer; font-weight: 600; font-size: 0.9rem; width: 100%; }
            .btn:hover { background: var(--accent-hover); }
            .status-box { font-size: 0.85rem; padding: 12px; border-radius: 6px; background: var(--bg); border-left: 4px solid var(--border); margin-top: 12px; word-break: break-all; }
            .chat-container { flex: 1; overflow-y: auto; padding: 30px; display: flex; flex-direction: column; gap: 20px; }
            .chat-bubble { max-width: 85%; padding: 16px 20px; border-radius: var(--radius); line-height: 1.6; font-size: 0.95rem; box-shadow: var(--shadow); white-space: pre-line; }
            .chat-bubble.user { background: var(--primary); color: white; align-self: flex-end; border-bottom-right-radius: 4px; }
            .chat-bubble.ai { background: white; color: var(--text-dark); align-self: flex-start; border-bottom-left-radius: 4px; border: 1px solid var(--border); }
            .input-area { background: var(--panel-bg); border-top: 1px solid var(--border); padding: 20px 30px; display: flex; gap: 12px; }
            .input-area input { flex: 1; padding: 14px 18px; border: 1px solid var(--border); border-radius: 30px; font-size: 0.95rem; outline: none; background: var(--bg); }
            .input-area input:focus { border-color: var(--accent); background: white; }
            .input-area button { width: 48px; height: 48px; border-radius: 50%; background: var(--accent); border: none; color: white; cursor: pointer; display: flex; align-items: center; justify-content: center; }
            @media (max-width: 768px) {
                .container { grid-template-columns: 1fr; }
                .sidebar { display: none; }
            }
        </style>
    </head>
    <body>
        <header>
            <h1>🏢 Enterprise Knowledge Hub <span class="badge">v2.5 Production</span></h1>
            <div style="font-size: 0.9rem; color: #94a3b8;" id="provider-badge">連線中...</div>
        </header>
        <div class="container">
            <div class="sidebar">
                <div class="card">
                    <h2>文件核心匯入區</h2>
                    <div class="file-dropzone" onclick="document.getElementById('fileInput').click()">
                        <p style="font-size: 0.85rem; font-weight: 500; color:var(--accent);">點擊此處選取內部檔案</p>
                        <p style="font-size: 0.75rem; color: var(--text-light); margin-top: 4px;">支援 PDF, DOCX, XLSX, TXT</p>
                    </div>
                    <input type="file" id="fileInput" onchange="document.getElementById('fn-disp').innerText='已選擇: '+this.files[0].name">
                    <div id="fn-disp" style="font-size:0.75rem; margin-top:6px; color:var(--accent);"></div>
                    <button class="btn" style="margin-top: 12px;" onclick="uploadFile()">開始結構化建檔</button>
                    <div id="uploadStatus" class="status-box" style="display: none;"></div>
                </div>
            </div>
            <div class="main-panel">
                <div class="chat-container" id="chat-box">
                    <div class="chat-bubble ai">
                        歡迎使用企業專屬知識庫！請先在左側面板上傳內部法規、合約或政策文件。上傳成功後，您可以在下方直接進行高精確度的流式 RAG 提問檢索。
                    </div>
                </div>
                <div class="input-area">
                    <input type="text" id="questionInput" placeholder="輸入企業規章、請假福利、流程規定等問題..." onkeypress="if(event.keyCode==13) sendQuery()">
                    <button onclick="sendQuery()">🚀</button>
                </div>
            </div>
        </div>
        <script>
            window.onload = async () => {
                const res = await fetch('/status'); const d = await res.json();
                document.getElementById('provider-badge').innerText = `核心驅動: ${d.provider.toUpperCase()} (${d.model})`;
            };
            async function uploadFile() {
                const fi = document.getElementById('fileInput'); const sb = document.getElementById('uploadStatus');
                if (!fi.files[0]) { alert('請選取檔案！'); return; }
                const fd = new FormData(); fd.append('file', fi.files[0]);
                sb.style.display = 'block'; sb.innerText = '系統進行深度清洗、切片與特徵工程中...';
                const response = await fetch('/upload', { method: 'POST', body: fd });
                const r = await response.json();
                if(response.ok) sb.innerText = `建檔成功！已切割 ${r.chunks_added} 個語義區塊並安全寫入磁碟。`;
                else sb.innerText = '錯誤: ' + r.detail;
            }
            function sendQuery() {
                const qi = document.getElementById('questionInput'); const cb = document.getElementById('chat-box');
                const q = qi.value.trim(); if (!q) return;
                cb.innerHTML += `<div class="chat-bubble user">${q}</div>`;
                const ab = document.createElement('div'); ab.className = 'chat-bubble ai'; ab.innerText = '檢索並生成中...';
                cb.appendChild(ab); cb.scrollTop = cb.scrollHeight; qi.value = '';
                const es = new EventSource(`/query?question=${encodeURIComponent(q)}&user=Guest`);
                let first = true;
                es.onmessage = function(e) {
                    if (first) { ab.innerText = ''; first = false; }
                    ab.innerText += e.data.replace(/\\\\n/g, '\\n'); cb.scrollTop = cb.scrollHeight;
                };
                es.onerror = function() { es.close(); };
            }
        </script>
    </body>
    </html>
    """

# ==========================================
# 4. 深度文件解析路由 (PDF / DOCX / XLSX / TXT)
# ==========================================
@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        filename = file.filename
        file_ext = os.path.splitext(filename)[1].lower()
        text_content = ""

        if file_ext == ".txt":
            text_content = contents.decode("utf-8", errors="ignore")
        elif file_ext == ".pdf":
            pdf_reader = pypdf.PdfReader(io.BytesIO(contents))
            text_content = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif file_ext == ".docx":
            doc = docx.Document(io.BytesIO(contents))
            text_content = "\n".join([p.text for p in doc.paragraphs])
        elif file_ext == ".xlsx":
            wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
            sheet_texts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = ", ".join([str(cell) for cell in row if cell is not None])
                    if row_text: sheet_texts.append(row_text)
            text_content = "\n".join(sheet_texts)
        else:
            raise HTTPException(status_code=400, detail=f"目前系統尚未支援 {file_ext} 格式檔案")

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="檔案解析成功但未讀取到任何有效文字內容")

        chunks_added = rag_engine.add_content(text=text_content, source_name=filename, source_type=file_ext)
        log_event("Document", f"解析並匯入文件: {filename}, 共 {chunks_added} 個片段")
        return {"status": "success", "filename": filename, "chunks_added": chunks_added}

    except Exception as e:
        log_event("Error", f"文件解析失敗: {str(e)}")
        raise HTTPException(status_code=500, detail=f"系統底層解析錯誤: {str(e)}")

@app.get("/query")
def chat_endpoint(question: str, user: str = "Guest"):
    return StreamingResponse(rag_engine.query_stream(question, user=user), media_type="text/event-stream")
